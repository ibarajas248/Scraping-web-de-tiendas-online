import pandas as pd, numpy as np, re, os
import matplotlib.pyplot as plt
from datetime import date
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.drawing.image import Image as XLImage
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =========================
# CONFIG
# =========================
INPUT_XLSX = "estudios_morrales.xlsx"

# Carpeta de salida (misma carpeta del script)
OUT_DIR = os.path.join(os.getcwd(), "salida_estudio_ultra_pro")
CHARTS_DIR = os.path.join(OUT_DIR, "charts_ultra_pro")
os.makedirs(CHARTS_DIR, exist_ok=True)

EXCEL_OUT = os.path.join(OUT_DIR, "ESTUDIO_MERCADO_MORRALES_ULTRA_PRO.xlsx")
WORD_OUT  = os.path.join(OUT_DIR, "ESTUDIO_MERCADO_MORRALES_ULTRA_PRO.docx")

# =========================
# LOAD
# =========================
df = pd.read_excel(INPUT_XLSX)

stores = [c for c in df.columns if c not in ['EAN', 'PRODUCTO', 'COINCIDENCIAS']]
for s in stores:
    df[s] = pd.to_numeric(df[s], errors="coerce")

df['EAN'] = df['EAN'].astype(str)

# =========================
# ENRICHMENT
# =========================
def segmento(text):
    t = str(text).lower()
    if re.search(r'escolar|colegio|niñ|kids|infant|disney|princes|mickey|minnie|hello kitty|barbie|pokemon|paw|spid|unicorn|sonic|jurassic|peppa|pep', t):
        return "Escolar / Kids"
    if re.search(r'laptop|notebook|porta ?comput|portacomp|pc\b|ejecutiv|oficina|business', t):
        return "Laptop / Oficina"
    if re.search(r'rued|trolley|arrastre|carro', t):
        return "Trolley / Ruedas"
    if re.search(r'outdoor|trek|hiking|camp|montañ|deport|sport|tactic|táctic|imperme|waterproof|viaje', t):
        return "Outdoor / Deportivo"
    return "Urbano / Casual"

df['segmento'] = df['PRODUCTO'].apply(segmento)

prices = df[stores]
df['n_precios'] = prices.notna().sum(axis=1)
df['precio_min'] = prices.min(axis=1, skipna=True)
df['precio_max'] = prices.max(axis=1, skipna=True)

# Precio de referencia: mínimo observado (proxy competitivo)
df['precio_ref'] = df['precio_min']

df['rango_abs'] = df['precio_max'] - df['precio_min']

# ✅ Evitar división por 0 o NaN
den = df['precio_ref'].replace({0: np.nan})
df['rango_pct'] = df['rango_abs'] / den

df['tienda_precio_min'] = prices.idxmin(axis=1, skipna=True)

bands = [0, 50_000, 80_000, 120_000, 160_000, 200_000, 250_000, 300_000, 500_000, 1_000_000, 10_000_000]
labels = ["<50k", "50–80k", "80–120k", "120–160k", "160–200k", "200–250k", "250–300k", "300–500k", "500k–1M", ">1M"]
df['banda_precio'] = pd.cut(df['precio_ref'], bins=bands, labels=labels, include_lowest=True, right=False)

# =========================
# KEY TABLES
# =========================
n_total = len(df)
n_priced = int(df['precio_ref'].notna().sum())
overlap_2p = int((df['n_precios'] >= 2).sum())

cobertura = df[stores].notna().sum().sort_values(ascending=False).reset_index()
cobertura.columns = ['Retailer', 'Productos_con_precio']
cobertura['%_del_dataset'] = (cobertura['Productos_con_precio'] / n_total).round(4)

band_table = df['banda_precio'].value_counts(dropna=False).reindex(labels, fill_value=0).reset_index()
band_table.columns = ['Banda_precio', 'Cantidad']
band_table['Participacion_%'] = (band_table['Cantidad'] / max(1, band_table['Cantidad'].sum()) * 100).round(2)

seg_summary = df.groupby('segmento').agg(
    Productos=('EAN', 'count'),
    Con_precio=('precio_ref', lambda x: x.notna().sum()),
    Mediana=('precio_ref', 'median'),
    P25=('precio_ref', lambda x: x.quantile(0.25)),
    P75=('precio_ref', lambda x: x.quantile(0.75)),
    Min=('precio_ref', 'min'),
    Max=('precio_ref', 'max'),
).reset_index().sort_values('Productos', ascending=False)

for c in ['Mediana', 'P25', 'P75', 'Min', 'Max']:
    seg_summary[c] = seg_summary[c].round(0)

store_stats = []
for s in stores:
    ser = df[s].dropna()
    store_stats.append([
        s, int(ser.count()),
        float(ser.median()) if len(ser) else np.nan,
        float(ser.quantile(0.25)) if len(ser) else np.nan,
        float(ser.quantile(0.75)) if len(ser) else np.nan,
        float(ser.min()) if len(ser) else np.nan,
        float(ser.max()) if len(ser) else np.nan
    ])

store_stats_df = pd.DataFrame(store_stats, columns=['Retailer', 'n', 'Mediana', 'P25', 'P75', 'Min', 'Max']).sort_values('n', ascending=False)
for c in ['Mediana', 'P25', 'P75', 'Min', 'Max']:
    store_stats_df[c] = store_stats_df[c].round(0)

# Segment share by retailer
store_seg = []
for s in stores:
    sub = df[df[s].notna()]
    vc = sub['segmento'].value_counts()
    tot = vc.sum()
    for seg, cnt in vc.items():
        store_seg.append([s, seg, int(cnt), float(cnt / tot) if tot else np.nan])

store_seg_df = pd.DataFrame(store_seg, columns=['Retailer', 'Segmento', 'n', 'share']).sort_values(['Retailer', 'n'], ascending=[True, False])

pivot_share = store_seg_df.pivot_table(index='Retailer', columns='Segmento', values='share', fill_value=0.0)
pivot_n = store_seg_df.pivot_table(index='Retailer', columns='Segmento', values='n', fill_value=0).astype(int)

# Overlap matrix
overlap = pd.DataFrame(index=stores, columns=stores, dtype=int)
for a in stores:
    for b in stores:
        overlap.loc[a, b] = int(((df[a].notna()) & (df[b].notna())).sum())
overlap_reset = overlap.reset_index().rename(columns={'index': 'Retailer'})

# Competitive price points by segment
def q(x, p):
    x = x.dropna()
    return float(x.quantile(p)) if len(x) else np.nan

targets = []
for seg in seg_summary['segmento']:
    x = df.loc[df['segmento'] == seg, 'precio_ref']
    targets.append([seg, int(x.notna().sum()), q(x, 0.4), q(x, 0.5), q(x, 0.6), q(x, 0.75)])

targets_df = pd.DataFrame(targets, columns=['Segmento', 'n_con_precio', 'P40', 'P50', 'P60', 'P75'])
for c in ['P40', 'P50', 'P60', 'P75']:
    targets_df[c] = targets_df[c].round(0)

line_arch = []
for _, r in targets_df.iterrows():
    if r['n_con_precio'] < 10:
        continue
    line_arch.append([r['Segmento'], r['P40'], r['P50'], r['P60'], r['P75']])

line_arch_df = pd.DataFrame(line_arch, columns=['Segmento', 'Entrada(P40)', 'Core_low(P50)', 'Core_high(P60)', 'Premium(P75)'])

# =========================
# CHARTS
# =========================
# A) Coverage bar
plt.figure()
cobertura.set_index('Retailer')['Productos_con_precio'].plot(kind='bar')
plt.title("Cobertura de productos con precio por retailer")
plt.ylabel("Cantidad de productos")
plt.tight_layout()
img_cov = os.path.join(CHARTS_DIR, "cobertura.png")
plt.savefig(img_cov)
plt.close()

# B) Price bands bar
plt.figure()
band_table.set_index('Banda_precio')['Cantidad'].plot(kind='bar')
plt.title("Distribución de precios (bandas COP)")
plt.ylabel("Cantidad de productos")
plt.tight_layout()
img_bands = os.path.join(CHARTS_DIR, "bandas.png")
plt.savefig(img_bands)
plt.close()

# C) Segment counts bar
plt.figure()
seg_summary.set_index('segmento')['Productos'].plot(kind='bar')
plt.title("Distribución por segmento de uso")
plt.ylabel("Cantidad de productos")
plt.tight_layout()
img_seg = os.path.join(CHARTS_DIR, "segmentos.png")
plt.savefig(img_seg)
plt.close()

# D) Median price by segment bar
plt.figure()
seg_summary.set_index('segmento')['Mediana'].sort_values().plot(kind='bar')
plt.title("Precio mediano por segmento (COP)")
plt.ylabel("Precio (COP)")
plt.tight_layout()
img_medseg = os.path.join(CHARTS_DIR, "mediana_segmento.png")
plt.savefig(img_medseg)
plt.close()

# E) Boxplot price by segment (✅ corregido: tick_labels)
plt.figure(figsize=(10, 6))
box_data = [df.loc[df['segmento'] == seg, 'precio_ref'].dropna().values for seg in seg_summary['segmento']]
plt.boxplot(
    box_data,
    tick_labels=seg_summary['segmento'],  # ✅ FIX Matplotlib 3.9+
    vert=True,
    showfliers=False
)
plt.title("Dispersión de precios por segmento (boxplot)")
plt.ylabel("Precio (COP)")
plt.xticks(rotation=20, ha='right')
plt.tight_layout()
img_box = os.path.join(CHARTS_DIR, "boxplot_segmento.png")
plt.savefig(img_box)
plt.close()

# F) Stacked share by retailer
plt.figure()
pivot_share.plot(kind='bar', stacked=True)
plt.title("Mix de segmentos por retailer (participación)")
plt.ylabel("Participación")
plt.legend(bbox_to_anchor=(1.02, 1), loc='upper left')
plt.tight_layout()
img_mix = os.path.join(CHARTS_DIR, "mix_retailer.png")
plt.savefig(img_mix)
plt.close()

# G) Overlap heatmap
plt.figure()
plt.imshow(overlap.values, aspect='auto')
plt.xticks(range(len(stores)), stores, rotation=45, ha='right')
plt.yticks(range(len(stores)), stores)
plt.title("Overlap (productos con precio en ambas tiendas)")
plt.colorbar()
plt.tight_layout()
img_overlap = os.path.join(CHARTS_DIR, "overlap.png")
plt.savefig(img_overlap)
plt.close()

# =========================
# EXCEL ULTRA PRO (tables + embedded charts)
# =========================
wb = Workbook()
wb.remove(wb.active)

def add_sheet(name, dframe, table=True):
    ws = wb.create_sheet(title=name[:31])
    for r in dataframe_to_rows(dframe, index=False, header=True):
        ws.append(r)

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    ws.freeze_panes = "A2"

    # Auto width
    for col in ws.columns:
        col_letter = col[0].column_letter
        max_len = 0
        for cell in col[:2000]:
            v = str(cell.value) if cell.value is not None else ""
            max_len = max(max_len, len(v))
        ws.column_dimensions[col_letter].width = min(65, max(10, max_len + 2))

    if table and ws.max_row > 1 and ws.max_column > 1:
        tab = Table(
            displayName=re.sub(r'[^A-Za-z0-9]', '', name)[:20] + "Tbl",
            ref=f"A1:{ws.cell(row=ws.max_row, column=ws.max_column).coordinate}"
        )
        tab.tableStyleInfo = TableStyleInfo(name="TableStyleMedium9", showRowStripes=True)
        ws.add_table(tab)

    return ws

add_sheet("00_Resumen", pd.DataFrame({
    'Metrica': ['Productos (total)', 'Productos con precio', 'Retailers', 'Productos con 2+ precios'],
    'Valor': [n_total, n_priced, len(stores), overlap_2p],
    'Nota': ['Base completa', 'Filas con al menos 1 precio', 'Columnas de tiendas', 'Subconjunto comparable']
}), table=True)

add_sheet("01_Cobertura", cobertura, table=True)
add_sheet("02_Bandas_precio", band_table, table=True)
add_sheet("03_Segmentos", seg_summary, table=True)
add_sheet("04_Posicionamiento", store_stats_df, table=True)
add_sheet("05_Mix_seg_por_tienda_n", pivot_n.reset_index(), table=True)
add_sheet("06_Mix_seg_por_tienda_%", pivot_share.reset_index(), table=True)
add_sheet("07_Overlap", overlap_reset, table=True)
add_sheet("08_Puntos_precio", targets_df, table=True)
add_sheet("09_Arquitectura_linea", line_arch_df, table=True)

base_cols = ['EAN','PRODUCTO','segmento','precio_ref','banda_precio','n_precios','precio_min','precio_max','rango_abs','rango_pct','tienda_precio_min'] + stores + ['COINCIDENCIAS']
add_sheet("10_Base_enriquecida", df[base_cols].copy(), table=True)

# Charts sheet
ws_ch = wb.create_sheet("Graficos")
ws_ch['A1'] = "Gráficos (imágenes)"
ws_ch['A1'].font = Font(bold=True, size=14)

img_positions = [
    ("A3", img_cov), ("A22", img_bands), ("A41", img_seg), ("A60", img_medseg),
    ("J3", img_box), ("J28", img_mix), ("J58", img_overlap)
]
for cell, pth in img_positions:
    im = XLImage(pth)
    im.width = int(im.width * 0.9)
    im.height = int(im.height * 0.9)
    ws_ch.add_image(im, cell)

wb.save(EXCEL_OUT)

# =========================
# WORD ULTRA PRO (tables + charts)
# =========================
doc = Document()

h = doc.add_heading("ESTUDIO DE MERCADO DE MORRALES — ULTRA PRO (Fabricante)", 0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Base analizada: {INPUT_XLSX} | Fecha: {date.today().isoformat()}")

doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph(
    f"• Productos (EAN): {n_total}\n"
    f"• Filas con precio: {n_priced}\n"
    f"• Retailers/canales: {len(stores)} ({', '.join(stores)})\n"
    f"• Comparabilidad (2+ precios): {overlap_2p} productos (≈ {overlap_2p/n_total:.1%})\n\n"
    "Interpretación clave: el dataset permite un estudio robusto de rangos, escalera de precios y portafolios por canal. "
    "La comparación directa de precio entre tiendas es limitada por el bajo overlap de EAN."
)

doc.add_heading("Metodología (enfoque científico y de mercado)", level=1)
doc.add_paragraph(
    "1) Limpieza: conversión a numérico, manejo de nulos y estandarización.\n"
    "2) Precio de referencia (precio_ref): mínimo observado por producto entre retailers (proxy competitivo).\n"
    "3) Segmentación de mercado: clasificación semántica por reglas sobre el nombre del producto.\n"
    "4) Estadística descriptiva robusta: mediana y percentiles (P25–P75) por segmento y canal.\n"
    "5) Lectura estratégica: escalera de precios, mix de portafolio por canal y puntos de precio para arquitectura de línea."
)

doc.add_heading("1. Cobertura y representatividad por canal", level=1)
doc.add_paragraph("Cobertura = número de productos con precio detectado por retailer. Indica amplitud de portafolio observado.")
doc.add_picture(img_cov, width=Inches(6))

doc.add_paragraph("Tabla 1. Cobertura por retailer.")
t1 = doc.add_table(rows=1, cols=len(cobertura.columns))
for j, col in enumerate(cobertura.columns):
    t1.rows[0].cells[j].text = str(col)
for _, row in cobertura.iterrows():
    cells = t1.add_row().cells
    for j, col in enumerate(cobertura.columns):
        cells[j].text = str(row[col])

doc.add_heading("2. Escalera de precios del mercado", level=1)
doc.add_paragraph("Bandas de precio en COP: base para definir líneas entrada/core/premium.")
doc.add_picture(img_bands, width=Inches(6))

doc.add_paragraph("Tabla 2. Bandas de precio y participación.")
t2 = doc.add_table(rows=1, cols=len(band_table.columns))
for j, col in enumerate(band_table.columns):
    t2.rows[0].cells[j].text = str(col)
for _, row in band_table.iterrows():
    cells = t2.add_row().cells
    for j, col in enumerate(band_table.columns):
        cells[j].text = str(row[col])

doc.add_heading("3. Segmentación funcional (uso) y tamaño de mercado", level=1)
doc.add_paragraph("Distribución de productos por segmento, interpretada como mix de oferta observada.")
doc.add_picture(img_seg, width=Inches(6))

doc.add_paragraph("Tabla 3. Segmentos y rangos de precio (mediana y P25–P75).")
t3 = doc.add_table(rows=1, cols=len(seg_summary.columns))
for j, col in enumerate(seg_summary.columns):
    t3.rows[0].cells[j].text = str(col)
for _, row in seg_summary.iterrows():
    cells = t3.add_row().cells
    for j, col in enumerate(seg_summary.columns):
        cells[j].text = str(row[col])

doc.add_heading("4. Posicionamiento de precio por segmento", level=1)
doc.add_paragraph("Mediana por segmento + boxplot para dispersión: detecta segmentos naturalmente premium y su variabilidad.")
doc.add_picture(img_medseg, width=Inches(6))
doc.add_picture(img_box, width=Inches(6))

doc.add_heading("5. Mix de portafolio por canal (retailer)", level=1)
doc.add_paragraph("Participación por segmento dentro de cada retailer: ayuda a decidir dónde competir y con qué línea.")
doc.add_picture(img_mix, width=Inches(6))

doc.add_heading("6. Comparabilidad entre canales (overlap)", level=1)
doc.add_paragraph("Heatmap de productos que aparecen simultáneamente en pares de retailers. Bajo overlap limita benchmarks directos.")
doc.add_picture(img_overlap, width=Inches(6))

doc.add_heading("7. Puntos de precio objetivo y arquitectura de línea", level=1)
doc.add_paragraph(
    "Recomendación operativa para fabricación: fijar precios guía por percentiles.\n"
    "Entrada ≈ P40; Core ≈ P50–P60; Premium ≈ P75. Ajustar según costos, calidad, garantías y marca."
)

doc.add_paragraph("Tabla 4. Puntos de precio objetivo por segmento (P40/P50/P60/P75).")
t4 = doc.add_table(rows=1, cols=len(targets_df.columns))
for j, col in enumerate(targets_df.columns):
    t4.rows[0].cells[j].text = str(col)
for _, row in targets_df.iterrows():
    cells = t4.add_row().cells
    for j, col in enumerate(targets_df.columns):
        cells[j].text = str(row[col])

doc.add_paragraph("Tabla 5. Arquitectura de línea sugerida (entrada/core/premium).")
t5 = doc.add_table(rows=1, cols=len(line_arch_df.columns))
for j, col in enumerate(line_arch_df.columns):
    t5.rows[0].cells[j].text = str(col)
for _, row in line_arch_df.iterrows():
    cells = t5.add_row().cells
    for j, col in enumerate(line_arch_df.columns):
        cells[j].text = str(row[col])

doc.add_heading("8. Recomendaciones para completar el estudio (nivel industrial)", level=1)
doc.add_paragraph(
    "Para convertir este análisis en un estudio de mercado completo (y no solo de precio/portafolio), se recomienda capturar además:\n"
    "• Atributos técnicos: capacidad (L), material (poliéster/cordura/cuero), impermeabilidad, costuras, herrajes, peso.\n"
    "• Marca y línea/modelo (normalizados).\n"
    "• Señales de demanda: reviews, ranking, disponibilidad/stock, ventas (si se consigue), tiempo de entrega.\n"
    "• Fotografía/estética y claims (garantía, antifraude, reflectivos, USB, etc.).\n"
    "Con esto se puede estimar un modelo de 'precio hedónico' y cuantificar qué atributos justifican el premium."
)

doc.save(WORD_OUT)

print("OK ✅")
print("Excel:", EXCEL_OUT)
print("Word :", WORD_OUT)
